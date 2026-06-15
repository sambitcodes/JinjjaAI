import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    # Also print tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            fullText.append(" | ".join(row_text))
            
    return '\n'.join(fullText)

if __name__ == "__main__":
    content = read_docx(r"c:\Users\sambi\OneDrive\Desktop\Projects\HangeulAI\phase_updates\course0phase1.docx")
    print(content)

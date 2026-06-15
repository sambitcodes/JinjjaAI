import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

def read_docx(file_path):
    doc = docx.Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            fullText.append(" | ".join(row_text))
            
    return '\n'.join(fullText)

if __name__ == "__main__":
    import shutil
    import os
    # Copy file to avoid locks
    src = r"c:\Users\sambi\OneDrive\Desktop\Projects\HangeulAI\phase_updates\course0phase1.docx"
    dest = r"c:\Users\sambi\OneDrive\Desktop\Projects\HangeulAI\phase_updates\course0phase1_copy.docx"
    shutil.copyfile(src, dest)
    try:
        content = read_docx(dest)
        with open(r"c:\Users\sambi\OneDrive\Desktop\Projects\HangeulAI\brain\affd8385-82b6-479a-8fe5-ebf47ff26884\scratch\docx_content.txt", "w", encoding="utf-8") as f:
            f.write(content)
        print("Success")
    finally:
        if os.path.exists(dest):
            os.remove(dest)
